from docx import Document
import json
from sentence_transformers import SentenceTransformer
from sentence_transformers import util
model = SentenceTransformer("all-MiniLM-L6-v2")


doc = Document("document/docx/MSFT_FY25q4_10K.docx")
    


def extract_text_from_docx(file_path: str) -> tuple[list[str], list[list[list[str]]]]:
    doc = Document(file_path) #read from the paragra
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # the strip means get all the space out 
    # for each line of paragrap add into paras in to a list 

    tables=[]
    for t in doc.tables:
        #this will get all the table 
        rows=[] #一行
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
            # this was each cell content add into each sell in rows
        tables.append(rows)
        # this will add all the rows into tables
    return paras, tables

def copyExistIntoFileTest(para: list,tables:list,file_path:str)->bool:
    try:
        with open (file_path,"w",encoding="utf-8") as f:
            for p in para:
                f.write(p +"\n") # this will only write the paragraph into the file without table
            f.write("\n\n=== TABLES ===\n")# this is write the table title into the file
            for ti,table in enumerate(tables,1):
                f.write(f"\n[TABLE {ti}]\n") # this will write the table number into the file
                for row in table:
                    f.write("\t".join(row) + "\n") # this will write the table content into the file with tab space between each cell
        return True
    except Exception as e:
        print(f"Error writing to file: {e}")
        return False
    
# this is read from the output again to chunk them   // document only can read docx
def get_chunk_paragraphs(txt_path: str) -> list[str]:
    with open(txt_path, "r", encoding="utf-8") as f:
        return [line.strip() for line in f if line.strip()]
def chunk_paragraphs(paragra:list[str],target_chars:int = 2500)->list[str]:
    chunks=[]
    buf=[]
    n=0
    for p in paragra:
       if n+ len(p) >target_chars and buf:
           chunks.append("\n".join(buf))
           buf=[]
           n=0
       buf.append(p)
       n+=len(p)
    if buf:
        chunks.append("\n".join(buf))
    return chunks
def read_chunks_jsonl(jsonl_path: str) -> list[dict]:
    chunks = []
    with open(jsonl_path, "r", encoding="utf-8") as f:
        for line in f:
            chunks.append(json.loads(line))
    return chunks



def embed_chunks(chunks: list[dict]) -> list[dict]:
    texts = [chunk["text"] for chunk in chunks]
    embeddings = model.encode(texts, normalize_embeddings=True)
    return embeddings
def write_chunks_jsonl(chunks: list[str], out_path: str) -> None:
    with open(out_path, "w", encoding="utf-8") as f:
        for i, text in enumerate(chunks):
            rec = {"chunk_id": i, "text": text}
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")

def main():
    outputPara=[]
    chunkFromjsonl=[]
    file_path = "document/docx/MSFT_FY25q4_10K.docx"
    output_path = "output.txt"
    paras, tables = extract_text_from_docx(file_path)
    success = copyExistIntoFileTest(paras, tables, output_path)
    if success:
        outputPara = get_chunk_paragraphs(output_path)
        outputPara = chunk_paragraphs(outputPara)   
        write_chunks_jsonl(outputPara, "output.jsonl")
        chunkFromjsonl=read_chunks_jsonl("output.jsonl")
        jsonAfterEmbed=embed_chunks(chunkFromjsonl)
        q = "how much money did Microsoft make in this?"
        q_vec = model.encode([q], normalize_embeddings=True)
        hits = util.semantic_search(q_vec, jsonAfterEmbed, top_k=5)[0]
        for h in hits:
            idx = h["corpus_id"]
            print("score=", h["score"], "chunk_id=", chunkFromjsonl[idx]["chunk_id"])
            print(chunkFromjsonl[idx]["text"][:200], "\n---\n")
    else:
        print("Failed to write content to file.")   



if __name__ == "__main__":
        main()